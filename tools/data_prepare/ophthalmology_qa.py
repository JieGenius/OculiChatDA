import json
import random

from docx import Document
import os
import os.path as osp


def parse_document(doc_path):
    doc = Document(doc_path)
    data = {}

    current_question = None
    current_idx = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Answers"):
            break  # Stop parsing when answers section is reached

        if paragraph.text.strip().endswith(":") or paragraph.text.strip().endswith("?"):
            # New question found
            if current_question:
                data[current_idx] = current_question
            current_idx, question = paragraph.text.strip().split(")", maxsplit=1)
            current_question = {"Question": question, "options": []}

        elif current_question:
            # Collect options for the current question
            if paragraph.text.strip():
                current_question["options"].append(paragraph.text.strip())

    if current_question:
        data[current_idx] = current_question

    return data


def parse_answers(doc_path):
    doc = Document(doc_path)
    answers = {}

    reading_answers = False

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Answers"):
            reading_answers = True
        elif reading_answers and paragraph.text.strip():
            # Assuming format "X. Answer"
            question_number, answer = paragraph.text.strip().split(".", maxsplit=1)
            answers[question_number] = answer

    return answers


def format_data(questions, answers):
    formatted_data = []

    for num, question in questions.items():
        formatted_question = {
            "Question": question["Question"],
            "options": question["options"],
            "Answer": answers.get(num, "")
        }
        formatted_data.append(formatted_question)

    return formatted_data


def generate_llm_data(formatted_data):
    data = []
    for item in formatted_data:
        question = item["Question"].strip()
        options = item["options"]
        answer = item["Answer"]
        data.append({
            "conversation": [
                {
                    "system": "你是一名眼科专家，你需要解答患者的疑问，提供准确的回答，必要时，提醒患者及时挂号就医。",
                    "input": question + "\n" + "\n".join(
                        [chr((ord("A") + i)) + "." + o for i, o in enumerate(options)]),
                    "output":( "正确答案为：" + answer) + ("" if random.randint(0, 1) == 0 else "\n你还有什么其他问题吗？")
                }
            ]
        })
    # convert ascill to char

    return data


if __name__ == "__main__":
    doc_path = "../../data/raw_data/Ophthalmology_Questions_and_Answers.docx"

    questions = parse_document(doc_path)
    answers = parse_answers(doc_path)
    formatted_data = format_data(questions, answers)

    # for item in formatted_data:
    #     print(item)

    res = generate_llm_data(formatted_data)
    with open("../../data/processed_data/ophthalmology_qa.json", "w") as f:
        json.dump(res, f, ensure_ascii=False, indent=4)